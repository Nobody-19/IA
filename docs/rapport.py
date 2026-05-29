# docs/rapport.py
import os
import json
from pathlib import Path
from datetime import datetime
from dotenv import load_dotenv
from groq import Groq

# Export Word
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Export PDF
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable

BASE_DIR = Path(__file__).parent.parent
ENV_FILE = BASE_DIR / ".env"
load_dotenv(dotenv_path=ENV_FILE)

GROQ_MODEL = "llama-3.3-70b-versatile"

# ──────────────────────────────────────────
# PROMPTS
# ──────────────────────────────────────────

PROMPT_RESUME = """Tu es un assistant de direction expert en synthèse de documents d'entreprise.

Génère un RÉSUMÉ COURT (1 page maximum) du document fourni.
Réponds UNIQUEMENT en JSON valide avec cette structure exacte :

{{
  "titre": "Résumé — [nom du document]",
  "date_rapport": "{date}",
  "type": "resume",
  "points_cles": [
    "Point clé 1",
    "Point clé 2",
    "Point clé 3",
    "Point clé 4",
    "Point clé 5"
  ],
  "chiffres_importants": [
    {{"label": "Chiffre 1", "valeur": "..."}},
    {{"label": "Chiffre 2", "valeur": "..."}}
  ],
  "conclusion": "Conclusion en 2-3 phrases maximum."
}}

Contenu du document :
{contexte}

Demande : {question}"""

PROMPT_BRIEFING = """Tu es un assistant de direction expert en analyse de documents d'entreprise.

Génère un BRIEFING DÉTAILLÉ du document fourni pour le Directeur Général.
Réponds UNIQUEMENT en JSON valide avec cette structure exacte :

{{
  "titre": "Briefing — [nom du document]",
  "date_rapport": "{date}",
  "type": "briefing",
  "resume_executif": "Paragraphe de 3-4 phrases résumant l'essentiel.",
  "sections": [
    {{
      "titre": "1. Contexte et objectifs",
      "contenu": "..."
    }},
    {{
      "titre": "2. Points financiers / chiffres clés",
      "contenu": "..."
    }},
    {{
      "titre": "3. Analyse et observations",
      "contenu": "..."
    }},
    {{
      "titre": "4. Points d'attention",
      "contenu": "..."
    }}
  ],
  "chiffres_importants": [
    {{"label": "Indicateur", "valeur": "Valeur"}},
    {{"label": "Indicateur 2", "valeur": "Valeur 2"}}
  ],
  "recommandations": [
    "Recommandation 1",
    "Recommandation 2",
    "Recommandation 3"
  ],
  "conclusion": "Conclusion générale en 2-3 phrases."
}}

Contenu du document :
{contexte}

Demande : {question}"""

# ──────────────────────────────────────────
# DÉTECTION DU TYPE DE RAPPORT
# ──────────────────────────────────────────
def detecter_type_rapport(question: str) -> str:
    """Détecte si l'utilisateur veut un résumé court ou un briefing détaillé."""
    mots_resume = ["résumé", "resume", "court", "bref", "rapide", "synthèse", "synthese", "essentiel"]
    mots_briefing = ["briefing", "détaillé", "detaille", "complet", "analyse", "rapport", "breffing", "approfondi"]

    q = question.lower()
    score_briefing = sum(1 for m in mots_briefing if m in q)
    score_resume   = sum(1 for m in mots_resume   if m in q)

    return "briefing" if score_briefing >= score_resume else "resume"


# ──────────────────────────────────────────
# GÉNÉRATION DU RAPPORT VIA GROQ
# ──────────────────────────────────────────
def generer_rapport(contexte: str, question: str, type_rapport: str = None) -> dict:
    """Génère le rapport structuré via Groq."""
    client = Groq(api_key=os.environ["GROQ_API_KEY"])
    date   = datetime.now().strftime("%d/%m/%Y")

    if type_rapport is None:
        type_rapport = detecter_type_rapport(question)

    prompt = (PROMPT_BRIEFING if type_rapport == "briefing" else PROMPT_RESUME).format(
        contexte=contexte[:6000],  # limite tokens
        question=question,
        date=date
    )

    response = client.chat.completions.create(
        model=GROQ_MODEL,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.2,
        max_tokens=2048,
    )

    text = response.choices[0].message.content.strip()
    text = text.replace("```json", "").replace("```", "").strip()

    try:
        return json.loads(text)
    except json.JSONDecodeError:
        # Fallback structuré
        return {
            "titre": f"Rapport — {datetime.now().strftime('%d/%m/%Y')}",
            "date_rapport": date,
            "type": type_rapport,
            "resume_executif": text,
            "sections": [],
            "points_cles": [text],
            "chiffres_importants": [],
            "recommandations": [],
            "conclusion": ""
        }


# ──────────────────────────────────────────
# EXPORT WORD (.docx)
# ──────────────────────────────────────────
def exporter_word(rapport: dict, output_path: Path) -> Path:
    """Génère un fichier Word professionnel."""
    doc = DocxDocument()

    # Style général
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ── En-tête ──
    titre_par = doc.add_heading(rapport.get("titre", "Rapport"), level=0)
    titre_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titre_run = titre_par.runs[0]
    titre_run.font.color.rgb = RGBColor(0x1A, 0x56, 0x76)

    # Date
    date_par = doc.add_paragraph(f"Date : {rapport.get('date_rapport', '')}")
    date_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_par.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph()

    # ── Résumé court ──
    if rapport.get("type") == "resume":
        doc.add_heading("Points clés", level=1)
        for point in rapport.get("points_cles", []):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(point)

        if rapport.get("chiffres_importants"):
            doc.add_heading("Chiffres importants", level=1)
            for c in rapport["chiffres_importants"]:
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(f"{c.get('label', '')} : ")
                run.bold = True
                p.add_run(c.get('valeur', ''))

        if rapport.get("conclusion"):
            doc.add_heading("Conclusion", level=1)
            doc.add_paragraph(rapport["conclusion"])

    # ── Briefing détaillé ──
    else:
        if rapport.get("resume_executif"):
            doc.add_heading("Résumé exécutif", level=1)
            doc.add_paragraph(rapport["resume_executif"])

        for section in rapport.get("sections", []):
            doc.add_heading(section.get("titre", ""), level=2)
            doc.add_paragraph(section.get("contenu", ""))

        if rapport.get("chiffres_importants"):
            doc.add_heading("Indicateurs clés", level=1)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = "Indicateur"
            hdr[1].text = "Valeur"
            for cell in hdr:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for c in rapport["chiffres_importants"]:
                row = table.add_row().cells
                row[0].text = c.get("label", "")
                row[1].text = c.get("valeur", "")

        if rapport.get("recommandations"):
            doc.add_heading("Recommandations", level=1)
            for rec in rapport["recommandations"]:
                p = doc.add_paragraph(style='List Number')
                p.add_run(rec)

        if rapport.get("conclusion"):
            doc.add_heading("Conclusion", level=1)
            doc.add_paragraph(rapport["conclusion"])

    # Pied de page
    section_doc = doc.sections[0]
    footer = section_doc.footer
    footer_par = footer.paragraphs[0]
    footer_par.text = f"Document généré par Agent IA — {rapport.get('date_rapport', '')}"
    footer_par.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(str(output_path))
    return output_path


# ──────────────────────────────────────────
# EXPORT PDF
# ──────────────────────────────────────────
def exporter_pdf(rapport: dict, output_path: Path) -> Path:
    """Génère un fichier PDF professionnel."""
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        rightMargin=2*cm, leftMargin=2*cm,
        topMargin=2.5*cm, bottomMargin=2*cm
    )

    styles = getSampleStyleSheet()
    couleur_primaire = colors.HexColor("#1A5676")

    style_titre = ParagraphStyle(
        'Titre', parent=styles['Title'],
        fontSize=20, textColor=couleur_primaire,
        spaceAfter=6, alignment=1
    )
    style_h1 = ParagraphStyle(
        'H1', parent=styles['Heading1'],
        fontSize=14, textColor=couleur_primaire,
        spaceBefore=14, spaceAfter=6
    )
    style_h2 = ParagraphStyle(
        'H2', parent=styles['Heading2'],
        fontSize=12, textColor=colors.HexColor("#2E7D9B"),
        spaceBefore=10, spaceAfter=4
    )
    style_body = ParagraphStyle(
        'Body', parent=styles['Normal'],
        fontSize=11, leading=16, spaceAfter=6
    )
    style_bullet = ParagraphStyle(
        'Bullet', parent=styles['Normal'],
        fontSize=11, leading=16,
        leftIndent=20, spaceAfter=4,
        bulletIndent=10
    )
    style_date = ParagraphStyle(
        'Date', parent=styles['Normal'],
        fontSize=10, textColor=colors.grey,
        alignment=1, spaceAfter=12
    )

    story = []

    # Titre
    story.append(Paragraph(rapport.get("titre", "Rapport"), style_titre))
    story.append(Paragraph(f"Date : {rapport.get('date_rapport', '')}", style_date))
    story.append(HRFlowable(width="100%", thickness=1, color=couleur_primaire))
    story.append(Spacer(1, 0.4*cm))

    # ── Résumé court ──
    if rapport.get("type") == "resume":
        story.append(Paragraph("Points clés", style_h1))
        for point in rapport.get("points_cles", []):
            story.append(Paragraph(f"• {point}", style_bullet))

        if rapport.get("chiffres_importants"):
            story.append(Spacer(1, 0.3*cm))
            story.append(Paragraph("Chiffres importants", style_h1))
            data = [["Indicateur", "Valeur"]] + [
                [c.get("label",""), c.get("valeur","")]
                for c in rapport["chiffres_importants"]
            ]
            t = Table(data, colWidths=[10*cm, 6*cm])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), couleur_primaire),
                ('TEXTCOLOR',  (0,0), (-1,0), colors.white),
                ('FONTNAME',   (0,0), (-1,0), 'Helvetica-Bold'),
                ('FONTSIZE',   (0,0), (-1,-1), 10),
                ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor("#EAF4F9")]),
                ('GRID',       (0,0), (-1,-1), 0.5, colors.lightgrey),
                ('TOPPADDING', (0,0), (-1,-1), 5),
                ('BOTTOMPADDING', (0,0), (-1,-1), 5),
            ]))
            story.append(t)

        if rapport.get("conclusion"):
            story.append(Spacer(1, 0.3*cm))
            story.append(Paragraph("Conclusion", style_h1))
            story.append(Paragraph(rapport["conclusion"], style_body))

    # ── Briefing détaillé ──
    else:
        if rapport.get("resume_executif"):
            story.append(Paragraph("Résumé exécutif", style_h1))
            story.append(Paragraph(rapport["resume_executif"], style_body))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey))

        for section in rapport.get("sections", []):
            story.append(Paragraph(section.get("titre",""), style_h2))
            story.append(Paragraph(section.get("contenu",""), style_body))

        if rapport.get("chiffres_importants"):
            story.append(Paragraph("Indicateurs clés", style_h1))
            data = [["Indicateur", "Valeur"]] + [
                [c.get("label",""), c.get("valeur","")]
                for c in rapport["chiffres_importants"]
            ]
            t = Table(data, colWidths=[10*cm, 6*cm])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), couleur_primaire),
                ('TEXTCOLOR',  (0,0), (-1,0), colors.white),
                ('FONTNAME',   (0,0), (-1,0), 'Helvetica-Bold'),
                ('FONTSIZE',   (0,0), (-1,-1), 10),
                ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor("#EAF4F9")]),
                ('GRID',       (0,0), (-1,-1), 0.5, colors.lightgrey),
                ('TOPPADDING', (0,0), (-1,-1), 5),
                ('BOTTOMPADDING', (0,0), (-1,-1), 5),
            ]))
            story.append(t)

        if rapport.get("recommandations"):
            story.append(Spacer(1, 0.3*cm))
            story.append(Paragraph("Recommandations", style_h1))
            for i, rec in enumerate(rapport["recommandations"], 1):
                story.append(Paragraph(f"{i}. {rec}", style_bullet))

        if rapport.get("conclusion"):
            story.append(Spacer(1, 0.3*cm))
            story.append(HRFlowable(width="100%", thickness=1, color=couleur_primaire))
            story.append(Paragraph("Conclusion", style_h1))
            story.append(Paragraph(rapport["conclusion"], style_body))

    doc.build(story)
    return output_path


# ──────────────────────────────────────────
# FONCTION PRINCIPALE
# ──────────────────────────────────────────
def creer_rapport(contexte: str, question: str, output_dir: Path, type_rapport: str = None) -> dict:
    """
    Pipeline complet :
    1. Génère le rapport structuré
    2. Exporte en PDF et Word
    3. Retourne les chemins + contenu
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    date_str   = datetime.now().strftime("%Y%m%d_%H%M%S")
    type_final = type_rapport or detecter_type_rapport(question)

    # 1. Générer
    rapport = generer_rapport(contexte, question, type_final)

    # 2. Exporter
    nom_base   = f"rapport_{type_final}_{date_str}"
    path_pdf   = output_dir / f"{nom_base}.pdf"
    path_docx  = output_dir / f"{nom_base}.docx"

    exporter_pdf(rapport, path_pdf)
    exporter_word(rapport, path_docx)

    return {
        "rapport":    rapport,
        "path_pdf":   path_pdf,
        "path_docx":  path_docx,
        "type":       type_final,
    }